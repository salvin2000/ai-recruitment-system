"""
Zecpath – Resume Text Extraction Engine
Day 5 Deliverable
-----------------
Converts PDF and DOCX resumes into clean, structured text
ready for AI/NLP processing by the ATS module.
"""

import os
import re
import json
import logging
from pathlib import Path
from datetime import datetime

# PDF parsing
import pdfplumber

# DOCX parsing
from docx import Document as DocxDocument

# ─────────────────────────── Logging Setup ────────────────────────────────

logging.basicConfig(
    level=logging.INFO,
    format="%(asctime)s  [%(levelname)s]  %(message)s",
    datefmt="%H:%M:%S"
)
log = logging.getLogger("zecpath.extractor")

# ─────────────────────────── Constants ────────────────────────────────────

SECTION_KEYWORDS = [
    "education", "experience", "work experience", "skills", "technical skills",
    "projects", "certifications", "achievements", "awards", "languages",
    "summary", "objective", "profile", "interests", "references",
    "publications", "volunteering", "internship", "training"
]

NOISE_PATTERNS = [
    r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]",   # control chars
    r"[│┃|]{2,}",                              # table borders
    r"[─━═\-=_*]{3,}",                          # decorative dividers (ASCII + Unicode box-drawing)
    r"\s{3,}",                                  # excessive whitespace (replaced with 2 spaces)
]

# ─────────────────────────── Text Cleaning ────────────────────────────────

def _clean_line(line: str) -> str:
    """Apply noise removal and normalization to a single line."""
    # Remove control characters
    line = re.sub(r"[\x00-\x08\x0b\x0c\x0e-\x1f\x7f]", "", line)

    # Remove decorative dividers (box-drawing lines, ASCII rule lines)
    line = re.sub(r"[─━═\-=_*]{3,}", "", line)

    # Replace smart quotes, em-dashes etc. with plain ASCII equivalents
    replacements = {
        "\u2019": "'", "\u2018": "'",
        "\u201c": '"', "\u201d": '"',
        "\u2013": "-", "\u2014": "-",
        "\u2022": "-", "\u25cf": "-",
        "\u00b7": "-", "\u2023": "-",
        "\u25aa": "-", "\u27a2": "-",
    }
    for char, replacement in replacements.items():
        line = line.replace(char, replacement)

    # Collapse runs of spaces/tabs (but keep indent intent)
    line = re.sub(r"\t", "  ", line)
    line = re.sub(r" {3,}", "  ", line)

    # Strip trailing whitespace
    line = line.rstrip()
    return line


def _normalize_section_heading(line: str) -> str:
    """
    Detect and normalise section headings.
    Headings are returned in UPPER CASE; non-headings unchanged.
    """
    stripped = line.strip().lower()
    # Remove trailing colons, bullets, dashes
    cleaned = re.sub(r"[:\-\|]+$", "", stripped).strip()

    if cleaned in SECTION_KEYWORDS:
        return line.strip().upper()
    return line


def _normalize_bullet(line: str) -> str:
    """Standardise all bullet-point variants to a single dash."""
    return re.sub(r"^\s*[\u2022\u2023\u25cf\u25aa\u27a2\u00b7\*\+\-]\s+", "- ", line)


def clean_text(raw: str) -> str:
    """
    Full cleaning pipeline applied to raw extracted text.
    Steps:
      1. Split into lines
      2. Per-line noise removal
      3. Section heading normalisation
      4. Bullet normalisation
      5. Remove blank-line clusters (max 1 consecutive blank)
      6. Strip leading/trailing whitespace from whole document
    """
    lines = raw.splitlines()
    cleaned = []
    prev_blank = False

    for line in lines:
        line = _clean_line(line)
        line = _normalize_section_heading(line)
        line = _normalize_bullet(line)

        is_blank = len(line.strip()) == 0
        if is_blank:
            if not prev_blank:
                cleaned.append("")
            prev_blank = True
        else:
            cleaned.append(line)
            prev_blank = False

    return "\n".join(cleaned).strip()


# ─────────────────────────── PDF Reader ───────────────────────────────────

def extract_pdf(filepath: str) -> dict:
    """
    Extract text from a PDF resume.
    Handles: single-column, multi-column, tables.
    Returns a result dict.
    """
    result = {
        "source": filepath,
        "format": "PDF",
        "pages": 0,
        "raw_length": 0,
        "clean_length": 0,
        "text": "",
        "warnings": []
    }

    try:
        raw_pages = []
        with pdfplumber.open(filepath) as pdf:
            result["pages"] = len(pdf.pages)
            for i, page in enumerate(pdf.pages):
                page_text = page.extract_text(x_tolerance=3, y_tolerance=3)

                # Fallback: try extracting words if text is None or very short
                if not page_text or len(page_text.strip()) < 20:
                    words = page.extract_words()
                    if words:
                        page_text = " ".join(w["text"] for w in words)
                        result["warnings"].append(
                            f"Page {i+1}: Used word-extraction fallback (may be image-based or multi-column)"
                        )
                    else:
                        result["warnings"].append(
                            f"Page {i+1}: No extractable text found — possible scanned image"
                        )

                if page_text:
                    raw_pages.append(page_text)

        raw = "\n\n".join(raw_pages)
        result["raw_length"] = len(raw)
        result["text"] = clean_text(raw)
        result["clean_length"] = len(result["text"])

    except Exception as e:
        result["warnings"].append(f"Extraction error: {str(e)}")
        log.error(f"PDF extraction failed for {filepath}: {e}")

    return result


# ─────────────────────────── DOCX Reader ──────────────────────────────────

def extract_docx(filepath: str) -> dict:
    """
    Extract text from a DOCX resume.
    Handles: paragraphs, headers, tables, text boxes (shapes).
    Returns a result dict.
    """
    result = {
        "source": filepath,
        "format": "DOCX",
        "pages": "N/A",
        "raw_length": 0,
        "clean_length": 0,
        "text": "",
        "warnings": []
    }

    try:
        doc = DocxDocument(filepath)
        segments = []

        # ── Main body paragraphs ──────────────────────────
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                segments.append(text)

        # ── Tables (cell-by-cell) ─────────────────────────
        for table in doc.tables:
            for row in table.rows:
                row_cells = []
                for cell in row.cells:
                    cell_text = cell.text.strip()
                    if cell_text:
                        row_cells.append(cell_text)
                if row_cells:
                    segments.append("  ".join(row_cells))

        # ── Headers & Footers ─────────────────────────────
        for section in doc.sections:
            for hdr_para in section.header.paragraphs:
                text = hdr_para.text.strip()
                if text:
                    segments.insert(0, text)  # put header at top

        raw = "\n".join(segments)
        result["raw_length"] = len(raw)
        result["text"] = clean_text(raw)
        result["clean_length"] = len(result["text"])

    except Exception as e:
        result["warnings"].append(f"Extraction error: {str(e)}")
        log.error(f"DOCX extraction failed for {filepath}: {e}")

    return result


# ─────────────────────────── Main Entry Point ─────────────────────────────

def extract_resume(filepath: str) -> dict:
    """
    Detect file type and dispatch to the correct extractor.
    Supported: .pdf, .docx, .doc (converted to docx)
    """
    path = Path(filepath)
    if not path.exists():
        return {"error": f"File not found: {filepath}"}

    ext = path.suffix.lower()
    log.info(f"Processing: {path.name}  [{ext.upper()}]")

    if ext == ".pdf":
        result = extract_pdf(str(path))
    elif ext in (".docx", ".doc"):
        result = extract_docx(str(path))
    else:
        return {"error": f"Unsupported file type: {ext}"}

    result["filename"] = path.name
    result["extracted_at"] = datetime.now().isoformat()
    return result


def extract_batch(input_dir: str, output_dir: str) -> list:
    """
    Process all PDF/DOCX files in input_dir.
    Saves cleaned text to output_dir/<filename>.txt
    Saves metadata JSON to output_dir/<filename>.json
    Returns a list of result dicts.
    """
    input_path = Path(input_dir)
    output_path = Path(output_dir)
    output_path.mkdir(parents=True, exist_ok=True)

    results = []
    for file in sorted(input_path.glob("*")):
        if file.suffix.lower() not in (".pdf", ".docx", ".doc"):
            continue

        result = extract_resume(str(file))
        stem = file.stem

        # Save cleaned text
        txt_out = output_path / f"{stem}_cleaned.txt"
        txt_out.write_text(result.get("text", ""), encoding="utf-8")

        # Save metadata
        meta_out = output_path / f"{stem}_meta.json"
        meta = {k: v for k, v in result.items() if k != "text"}
        meta_out.write_text(json.dumps(meta, indent=2), encoding="utf-8")

        results.append(result)
        log.info(
            f"  -> Saved: {txt_out.name}  "
            f"({result.get('raw_length', 0)} raw chars → {result.get('clean_length', 0)} clean chars)"
        )

    return results


# ─────────────────────────── CLI ──────────────────────────────────────────

if __name__ == "__main__":
    import sys

    if len(sys.argv) < 2:
        print("Usage: python extractor.py <resume_file_or_directory> [output_dir]")
        sys.exit(1)

    target = sys.argv[1]
    out_dir = sys.argv[2] if len(sys.argv) > 2 else "output/cleaned"

    if Path(target).is_dir():
        results = extract_batch(target, out_dir)
        print(f"\nProcessed {len(results)} resume(s). Output in: {out_dir}")
    else:
        result = extract_resume(target)
        if "error" in result:
            print(f"Error: {result['error']}")
            sys.exit(1)
        out_path = Path(out_dir) / f"{Path(target).stem}_cleaned.txt"
        out_path.parent.mkdir(parents=True, exist_ok=True)
        out_path.write_text(result["text"], encoding="utf-8")
        print(f"Extracted: {result['clean_length']} characters → {out_path}")
